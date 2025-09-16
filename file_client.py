"""
file_client.py — wymiana plików asystent↔użytkownik + konwersje + OCR.
Sandbox w /workspace/mrd69. Integracje: memory, psychika, autonauka.
Obsługiwane: pdf, docx, pages, txt, md, jpg, png, zip, rar (opcjonalnie).
"""

from __future__ import annotations

import fnmatch
import hashlib
import json
import mimetypes
import os
import re
import shutil
import sys
import tempfile
import time
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any, cast

# ── Ścieżki
# Ustal ROOT względem pliku repo, a nie twardo zakodowanej ścieżki
REPO_ROOT = Path(__file__).parent.resolve()
ROOT = (REPO_ROOT / "data").resolve()
INBOX_DIR = (ROOT / "inbox").resolve()
OUT_DIR = (ROOT / "out" / "files").resolve()
PIPE_DIR = (ROOT / "out" / "writing" / "_pipe").resolve()
INDEX_JSON = OUT_DIR / "_index.json"
JOBS_FILE = PIPE_DIR / "jobs.jsonl"
for d in (INBOX_DIR, OUT_DIR, PIPE_DIR):
    d.mkdir(parents=True, exist_ok=True)

# ── Opcjonalne biblioteki
try:
    import pdfplumber  # lepszy ekstraktor PDF
except Exception:
    pdfplumber = None

try:
    import PyPDF2  # fallback PDF
except Exception:
    PyPDF2 = None

try:
    import docx  # python-docx
except Exception:
    docx = None

try:
    from PIL import Image, ImageOps
except Exception:
    Image = None  # type: ignore

try:
    import pytesseract  # OCR
except Exception:
    pytesseract = None  # type: ignore

try:
    import rarfile  # obsługa RAR (wymaga unrar/unarchiver w systemie)
except Exception:
    rarfile = None  # type: ignore

# ── Integracje
_HAS = {}


def _opt_import(name: str):
    try:
        m = __import__(name, fromlist=["*"])
        _HAS[name] = True
        return m
    except Exception:
        _HAS[name] = False
        return None


memory = _opt_import("memory")
psychika = _opt_import("psychika")
auto = _opt_import("autonauka")


# ── Utils
def now_ms() -> int:
    return int(time.time() * 1000)


def slug(s: str, max_len: int = 72) -> str:
    t = re.sub(r"\s+", "-", (s or "").strip())
    t = re.sub(r"[^a-zA-Z0-9\.\-_]+", "", t)
    t = t.strip("-").lower()
    return t[:max_len] or "file"


def ensure_parent(p: Path) -> None:
    p.parent.mkdir(parents=True, exist_ok=True)


def safe_join(base: Path, rel: str) -> Path:
    p = (base / rel).resolve()
    if not str(p).startswith(str(ROOT)):
        raise PermissionError("Path escapes ROOT")
    return p


def atomic_write_bytes(path: Path, data: bytes) -> None:
    ensure_parent(path)
    with tempfile.NamedTemporaryFile(dir=str(path.parent), delete=False) as tmp:
        tmp.write(data)
        tmp.flush()
        os.fsync(tmp.fileno())
        tpath = Path(tmp.name)
    tpath.replace(path)


def file_hash(path: Path, algo: str = "sha256", chunk: int = 65536) -> str:
    h = hashlib.new(algo)
    with path.open("rb") as f:
        while True:
            b = f.read(chunk)
            if not b:
                break
            h.update(b)
    return f"{algo}:{h.hexdigest()}"


def guess_ext_from_bytes(b: bytes, fallback: str = "") -> str:
    # heurystyki bez imghdr (usunięty w Python 3.13)
    text_like = b[:2048].decode("utf-8", errors="ignore")
    if re.search(r"%PDF-", text_like):
        return ".pdf"
    if re.search(r"\bPK\x03\x04", b[:4].decode("latin1", errors="ignore")):
        return ".zip"
    if text_like.count("\n") > 2:
        return ".txt"
    return fallback or ""


def detect_mime(path: Path) -> str:
    mt, _ = mimetypes.guess_type(str(path))
    if mt:
        return mt
    try:
        with path.open("rb") as f:
            head = f.read(12)
        if head.startswith(b"%PDF-"):
            return "application/pdf"
        # proste sygnatury obrazów
        if head.startswith(b"\xff\xd8\xff"):
            return "image/jpeg"
        if head.startswith(b"\x89PNG\r\n\x1a\n"):
            return "image/png"
        if head[:6] in (b"GIF87a", b"GIF89a"):
            return "image/gif"
        if head.startswith(b"II\x2a\x00") or head.startswith(b"MM\x00\x2a"):
            return "image/tiff"
        if head.startswith(b"RIFF") and b"WEBP" in head:
            return "image/webp"
    except Exception:
        pass
    return "application/octet-stream"


# ── Indeks + audit
def index_load() -> dict[str, Any]:
    if not INDEX_JSON.exists():
        return {"files": {}, "updated": now_ms()}
    try:
        return json.loads(INDEX_JSON.read_text(encoding="utf-8"))
    except Exception:
        return {"files": {}, "updated": now_ms()}


def index_save(idx: dict[str, Any]) -> None:
    idx["updated"] = now_ms()
    atomic_write_bytes(
        INDEX_JSON,
        json.dumps(idx, ensure_ascii=False, indent=2).encode("utf-8"),
    )


def touch_index(path: Path, kind: str, extra: dict[str, Any] | None = None) -> None:
    idx = index_load()
    rel = str(path.relative_to(ROOT)) if path.exists() else str(path)
    st = path.stat() if path.exists() else None
    info: dict[str, Any] = {
        "kind": kind,
        "mtime": int(st.st_mtime) if st else now_ms(),
        "size": int(st.st_size) if st else 0,
    }
    if extra:
        base_meta = cast(dict[str, Any], info.get("meta") or {})
        extra_dict = cast(dict[str, Any], extra)
        m: dict[str, Any] = {**base_meta, **extra_dict}
        info["meta"] = m
    idx["files"][rel] = info
    index_save(idx)


def record_job(event: str, payload: dict[str, Any]) -> None:
    ensure_parent(JOBS_FILE)
    with JOBS_FILE.open("a", encoding="utf-8") as f:
        f.write(json.dumps({"ts": now_ms(), "event": event, **payload}, ensure_ascii=False) + "\n")


# ── Integracje hooki
def mem_add(
    text: str,
    tags: list[str] | None = None,
    user: str = "global",
    conf: float = 0.7,
) -> None:
    if not memory:
        return
    try:
        if hasattr(memory, "add_fact"):
            memory.add_fact(
                text[:8000],
                tags=sorted(set(tags or ["files"])),
                user=user,
                conf=conf,
            )  # type: ignore
        elif hasattr(memory, "ltm_add_sync"):
            memory.ltm_add_sync(
                "[FILES] " + text[:120],
                sources=[{"text": text[:8000]}],
                user=user,
                tags=sorted(set(tags or ["files"])),
            )  # type: ignore
    except Exception:
        pass


def psy_event(kind: str, data: dict[str, Any]) -> None:
    if not _HAS.get("psychika"):
        return
    try:
        if hasattr(psychika, "psyche_event"):
            psychika.psyche_event(kind, data)  # type: ignore
        elif hasattr(psychika, "autopilot_cycle"):
            psychika.autopilot_cycle(  # pyright: ignore[reportOptionalMemberAccess]
                f"{kind}: {json.dumps(data, ensure_ascii=False)[:400]}"
            )  # type: ignore
    except Exception:
        pass


def auto_learn(sample: dict[str, Any]) -> None:
    if not _HAS.get("autonauka"):
        return
    try:
        for fn in ("learn", "enqueue", "add_sample"):
            if hasattr(auto, fn):
                getattr(auto, fn)(sample)  # type: ignore
                break
    except Exception:
        pass


# ── Wynik
@dataclass
class Result:
    ok: bool
    data: Any = None
    err: str = ""


# ── Upload/Download
def upload_file(data: bytes, filename: str, user: str = "global") -> Result:
    try:
        ext = Path(filename).suffix.lower()
        if not ext:
            ext = guess_ext_from_bytes(data, ".bin")
        name = f"{now_ms()}_{slug(Path(filename).name)}{ext}"
        p = INBOX_DIR / name
        atomic_write_bytes(p, data)
        hv = file_hash(p)
        touch_index(p, "upload", {"hash": hv, "filename": filename})
        mem_add(
            f"Upload: {filename} → {p} [{hv}]",
            tags=["files", "upload"],
            user=user,
            conf=0.75,
        )
        psy_event(
            "file_upload",
            {
                "path": str(p),
                "name": filename,
                "size": len(data),
            },
        )
        auto_learn(
            {
                "kind": "files.upload",
                "path": str(p),
                "name": filename,
                "size": len(data),
            }
        )
        record_job(
            "upload",
            {
                "path": str(p),
                "name": filename,
                "bytes": len(data),
                "hash": hv,
            },
        )
        return Result(True, {"path": str(p), "hash": hv, "bytes": len(data)})
    except Exception as e:
        psy_event("file_error", {"op": "upload", "name": filename, "err": str(e)})
        return Result(False, err=str(e))


def download_file(rel: str) -> Result:
    try:
        p = safe_join(ROOT, rel)
        if not p.exists() or not p.is_file():
            return Result(False, err="not_found")
        b = p.read_bytes()
        psy_event("file_download", {"path": str(p), "size": len(b)})
        record_job("download", {"path": str(p), "bytes": len(b)})
        return Result(True, {"bytes": b, "mime": detect_mime(p), "name": p.name})
    except Exception as e:
        return Result(False, err=str(e))


# ── Ekstrakcja tekstu
def pdf_to_text(path: Path) -> str:
    if pdfplumber:
        try:
            out = []
            with pdfplumber.open(str(path)) as pdf:
                for pg in pdf.pages:
                    out.append(pg.extract_text() or "")
            return "\n".join(out).strip()
        except Exception:
            pass
    if PyPDF2:
        try:
            reader = PyPDF2.PdfReader(str(path))
            return "\n".join([(p.extract_text() or "") for p in reader.pages]).strip()
        except Exception:
            pass
    return ""


def docx_to_text(path: Path) -> str:
    if not docx:
        return ""
    try:
        d = docx.Document(str(path))
        return "\n".join([p.text for p in d.paragraphs]).strip()
    except Exception:
        return ""


def pages_to_text(path: Path) -> str:
    # .pages to ZIP. Szukamy QuickLook/Preview.pdf lub index.xml
    try:
        with zipfile.ZipFile(str(path), "r") as z:
            # PDF preview
            for cand in (
                "QuickLook/Preview.pdf",
                "preview.pdf",
                "QuickLook/Thumbnail.jpg",
            ):
                if cand in z.namelist():
                    if cand.endswith(".pdf"):
                        with z.open(cand) as f:
                            tmp = Path(tempfile.mkstemp(suffix=".pdf")[1])
                            tmp.write_bytes(f.read())
                            txt = pdf_to_text(tmp)
                            try:
                                tmp.unlink()
                            except Exception:
                                pass
                            if txt:
                                return txt
                    elif cand.endswith(".jpg") and pytesseract and Image:
                        with z.open(cand) as f:
                            img_b = f.read()
                            tmp = Path(tempfile.mkstemp(suffix=".jpg")[1])
                            tmp.write_bytes(img_b)
                            t = ocr_image_to_text(tmp)
                            try:
                                tmp.unlink()
                            except Exception:
                                pass
                            if t:
                                return t
            # XML fallback
            for name in z.namelist():
                if name.lower().endswith(".xml"):
                    with z.open(name) as f:
                        xml = f.read().decode("utf-8", errors="ignore")
                        text = re.sub(r"<[^>]+>", " ", xml)
                        text = re.sub(r"\s+", " ", text)
                        return text.strip()
    except Exception:
        pass
    return ""


def image_to_text(path: Path) -> str:
    # OCR bezpieczny
    return ocr_image_to_text(path)


def ocr_image_to_text(path: Path) -> str:
    if not (pytesseract and Image):
        return ""
    try:
        img = Image.open(str(path))
        img = ImageOps.exif_transpose(img)
        # prosta normalizacja
        if max(img.size) > 2000:
            ratio = 2000.0 / max(img.size)
            img = img.resize((int(img.width * ratio), int(img.height * ratio)))
        img = img.convert("L")
        return pytesseract.image_to_string(img, lang="eng+pol").strip()
    except Exception:
        return ""


def extract_text(rel: str) -> Result:
    try:
        p = safe_join(ROOT, rel)
        if not p.exists():
            return Result(False, err="not_found")
        ext = p.suffix.lower()
        txt = ""
        if ext == ".pdf":
            txt = pdf_to_text(p)
        elif ext == ".docx":
            txt = docx_to_text(p)
        elif ext == ".pages":
            txt = pages_to_text(p)
        elif ext in (
            ".txt",
            ".md",
            ".json",
            ".csv",
            ".py",
            ".js",
            ".ts",
            ".html",
            ".css",
        ):
            try:
                txt = p.read_text(encoding="utf-8", errors="ignore")
            except Exception:
                txt = p.read_text(encoding="latin-1", errors="ignore")
        elif ext in (
            ".jpg",
            ".jpeg",
            ".png",
            ".bmp",
            ".tif",
            ".tiff",
            ".webp",
        ):
            txt = image_to_text(p)
        else:
            # zip/rar → listing treści
            if ext == ".zip":
                with zipfile.ZipFile(str(p), "r") as z:
                    txt = "\n".join(z.namelist())
            elif ext == ".rar" and rarfile:
                with rarfile.RarFile(str(p), "r") as z:  # type: ignore
                    txt = "\n".join(z.namelist())  # type: ignore
        psy_event("file_extract_text", {"path": str(p), "len": len(txt)})
        mem_add(f"Extract: {p.name}\n{txt[:800]}", tags=["files", "extract"])
        record_job("extract_text", {"path": str(p), "len": len(txt)})
        return Result(True, {"text": txt, "path": str(p)})
    except Exception as e:
        psy_event("file_error", {"op": "extract", "path": rel, "err": str(e)})
        return Result(False, err=str(e))


# ── Eksport/konwersje
def text_to_pdf(text: str, out_path: Path) -> bool:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.pdfgen import canvas
    except Exception:
        return False
    try:
        ensure_parent(out_path)
        c = canvas.Canvas(str(out_path), pagesize=A4)
        w, h = A4
        y = h - 2 * cm
        for line in text.splitlines():
            c.drawString(2 * cm, y, line[:110])
            y -= 14
            if y < 2 * cm:
                c.showPage()
                y = h - 2 * cm
        c.save()
        return True
    except Exception:
        return False


def text_to_docx(text: str, out_path: Path) -> bool:
    if not docx:
        return False
    try:
        ensure_parent(out_path)
        d = docx.Document()
        for para in text.split("\n\n"):
            d.add_paragraph(para)
        d.save(str(out_path))
        return True
    except Exception:
        return False


def make_zip(paths: list[str], out_rel: str) -> Result:
    try:
        out_path = safe_join(ROOT, out_rel)
        ensure_parent(out_path)
        with zipfile.ZipFile(str(out_path), "w", compression=zipfile.ZIP_DEFLATED) as z:
            for rp in paths:
                p = safe_join(ROOT, rp)
                if p.exists():
                    z.write(str(p), arcname=p.name)
        hv = file_hash(out_path)
        touch_index(out_path, "zip", {"hash": hv})
        mem_add(f"ZIP: {out_path.name} [{hv}]", tags=["files", "zip"])
        record_job("zip", {"out": str(out_path), "count": len(paths), "hash": hv})
        return Result(True, {"path": str(out_path), "hash": hv})
    except Exception as e:
        return Result(False, err=str(e))


def unzip_archive(rel: str, dest_rel: str) -> Result:
    try:
        arch = safe_join(ROOT, rel)
        dest = safe_join(ROOT, dest_rel)
        ensure_parent(dest / "x")
        if arch.suffix.lower() == ".zip":
            with zipfile.ZipFile(str(arch), "r") as z:
                z.extractall(str(dest))
        elif arch.suffix.lower() == ".rar" and rarfile:
            with rarfile.RarFile(str(arch), "r") as z:  # type: ignore
                z.extractall(str(dest))  # type: ignore
        else:
            return Result(False, err="unsupported_archive")
        record_job("unzip", {"arch": str(arch), "dest": str(dest)})
        psy_event("file_unzip", {"arch": str(arch), "dest": str(dest)})
        return Result(True, {"dest": str(dest)})
    except Exception as e:
        return Result(False, err=str(e))


def convert_file(rel: str, target: str, out_name: str | None = None) -> Result:
    """
    target: 'txt' | 'pdf' | 'docx' | 'zip' | 'png' | 'jpg'
    """
    try:
        src = safe_join(ROOT, rel)
        if not src.exists():
            return Result(False, err="not_found")
        target = target.lower().strip()
        stem = Path(out_name).stem if out_name else src.stem
        if target == "txt":
            txt = extract_text(rel)
            if not txt.ok:
                return txt
            out = OUT_DIR / f"{stem}.txt"
            atomic_write_bytes(out, txt.data["text"].encode("utf-8"))
        elif target == "pdf":
            txt = extract_text(rel)
            if not txt.ok:
                return txt
            out = OUT_DIR / f"{stem}.pdf"
            ok = text_to_pdf(txt.data["text"], out)
            if not ok:
                return Result(False, err="pdf_export_failed")
        elif target == "docx":
            txt = extract_text(rel)
            if not txt.ok:
                return txt
            out = OUT_DIR / f"{stem}.docx"
            ok = text_to_docx(txt.data["text"], out)
            if not ok:
                return Result(False, err="docx_export_failed")
        elif target == "zip":
            out = OUT_DIR / f"{stem}.zip"
            return make_zip([str(src.relative_to(ROOT))], str(out.relative_to(ROOT)))
        elif target in ("png", "jpg", "jpeg"):
            if not Image:
                return Result(False, err="pil_missing")
            img = Image.open(str(src))
            out = OUT_DIR / (f"{stem}.{ 'jpg' if target=='jpeg' else target }")
            img.save(str(out))
        else:
            return Result(False, err="unsupported_target")
        hv = file_hash(out)
        touch_index(out, "export", {"hash": hv, "src": str(src)})
        mem_add(f"Export: {src.name} → {out.name} [{hv}]", tags=["files", "export"])
        psy_event("file_export", {"src": str(src), "out": str(out)})
        record_job(
            "convert",
            {"src": str(src), "out": str(out), "target": target, "hash": hv},
        )
        return Result(True, {"path": str(out), "hash": hv})
    except Exception as e:
        psy_event("file_error", {"op": "convert", "src": rel, "err": str(e)})
        return Result(False, err=str(e))


# ── Proste I/O dla systemu (zachowane)
def list_dir(rel: str = "out/files", pattern: str = "*", recursive: bool = False) -> Result:
    try:
        base = safe_join(ROOT, rel)
        if not base.exists():
            return Result(True, [])
        out: list[str] = []
        if recursive:
            for p, _, fns in os.walk(base):
                for fn in fns:
                    fp = Path(p) / fn
                    if fnmatch.fnmatch(fn, pattern):
                        out.append(str(fp.relative_to(ROOT)))
        else:
            for fp in base.iterdir():
                if fp.is_file() and fnmatch.fnmatch(fp.name, pattern):
                    out.append(str(fp.relative_to(ROOT)))
        return Result(True, sorted(out))
    except Exception as e:
        return Result(False, err=str(e))


def read_text(rel: str, encoding: str = "utf-8") -> Result:
    try:
        p = safe_join(ROOT, rel)
        if not p.exists() or not p.is_file():
            return Result(False, err="not_found")
        return Result(True, p.read_text(encoding=encoding, errors="replace"))
    except Exception as e:
        return Result(False, err=str(e))


def write_text(rel: str, text: str, encoding: str = "utf-8") -> Result:
    try:
        p = safe_join(ROOT, rel)
        atomic_write_bytes(p, text.encode(encoding))
        hv = file_hash(p)
        touch_index(p, "file", {"hash": hv})
        mem_add(f"Write: {p} [{hv}]", tags=["files", "write"])
        psy_event("file_write", {"path": str(p), "size": len(text.encode(encoding))})
        record_job(
            "write",
            {"path": str(p), "bytes": len(text.encode(encoding)), "hash": hv},
        )
        return Result(True, {"path": str(p), "hash": hv})
    except Exception as e:
        return Result(False, err=str(e))


def delete_path(rel: str) -> Result:
    try:
        p = safe_join(ROOT, rel)
        if not p.exists():
            return Result(True, {"deleted": False})
        if p.is_dir():
            shutil.rmtree(p)
        else:
            p.unlink()
        touch_index(p, "deleted")
        mem_add(f"Delete: {p}", tags=["files", "delete"])
        psy_event("file_delete", {"path": str(p)})
        record_job("delete", {"path": str(p)})
        return Result(True, {"deleted": True})
    except Exception as e:
        return Result(False, err=str(e))


# ── CLI dla dev
def _print(obj: Any) -> None:
    print(json.dumps(obj, ensure_ascii=False, indent=2))


def _main(argv: list[str]) -> int:
    import argparse

    ap = argparse.ArgumentParser(description="file_client — upload/convert/OCR")
    sub = ap.add_subparsers(dest="cmd", required=True)

    up = sub.add_parser("upload")
    up.add_argument("--from-path", required=True)
    dl = sub.add_parser("download")
    dl.add_argument("--path", required=True)
    xt = sub.add_parser("extract")
    xt.add_argument("--path", required=True)
    cv = sub.add_parser("convert")
    cv.add_argument("--path", required=True)
    cv.add_argument("--to", required=True)
    cv.add_argument("--out-name", default="")
    ls = sub.add_parser("ls")
    ls.add_argument("--path", default="out/files")
    ls.add_argument("--pattern", default="*")
    ls.add_argument("--rec", action="store_true")
    rm = sub.add_parser("rm")
    rm.add_argument("--path", required=True)
    uz = sub.add_parser("unzip")
    uz.add_argument("--arch", required=True)
    uz.add_argument("--dest", required=True)
    zp = sub.add_parser("zip")
    zp.add_argument("--out", required=True)
    zp.add_argument("--paths", nargs="+", required=True)

    args = ap.parse_args(argv)

    if args.cmd == "upload":
        src = safe_join(ROOT, args.from_path)
        b = src.read_bytes()
        _print(upload_file(b, src.name).__dict__)
    elif args.cmd == "download":
        _print(download_file(args.path).__dict__)
    elif args.cmd == "extract":
        _print(extract_text(args.path).__dict__)
    elif args.cmd == "convert":
        _print(
            convert_file(
                args.path,
                args.to,
                out_name=args.out_name or "",
            ).__dict__
        )
    elif args.cmd == "ls":
        _print(list_dir(args.path, args.pattern, recursive=args.rec).__dict__)
    elif args.cmd == "rm":
        _print(delete_path(args.path).__dict__)
    elif args.cmd == "unzip":
        _print(unzip_archive(args.arch, args.dest).__dict__)
    elif args.cmd == "zip":
        _print(make_zip(args.paths, args.out).__dict__)
    return 0


if __name__ == "__main__":
    sys.exit(_main(sys.argv[1:]))
